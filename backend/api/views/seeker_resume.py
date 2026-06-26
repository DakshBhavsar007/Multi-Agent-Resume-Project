"""
Job Seeker Resume Views
────────────────────────
Handles:
  - Upload resume (PDF/DOCX) → parse with existing parsing_agent → store in JobSeekerAccount
  - Enhance resume with ResumeEnhancerAgent
  - Get current resume data
"""

import os
import uuid
import logging
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

from api.models import JobSeekerAccount
from api.views.seeker_auth import require_seeker_jwt
from api.services.email_service import FROM_EMAIL
from models.schemas import success_response, error_response
from agents.parsing_agent import ResumeParsingAgent
from agents.normalization_agent import SkillNormalizationAgent
from agents.resume_enhancer_agent import ResumeEnhancerAgent

logger = logging.getLogger(__name__)

UPLOAD_DIR = os.getenv("UPLOAD_DIR", "uploads")


@csrf_exempt
@require_seeker_jwt
def upload_resume(request):
    """
    POST /api/v1/seeker/resume/upload
    Upload a PDF/DOCX resume → parse → store in seeker account.
    """
    if request.method != "POST":
        return JsonResponse(error_response("Method not allowed"), status=405)
    try:
        file = request.FILES.get("file")
        if not file:
            return JsonResponse(error_response("No file provided"), status=400)

        allowed_ext = (".pdf", ".docx", ".doc", ".txt")
        if not file.name.lower().endswith(allowed_ext):
            return JsonResponse(error_response("Only PDF, DOCX, DOC, or TXT files are accepted"), status=400)

        if file.size > 10 * 1024 * 1024:  # 10 MB
            return JsonResponse(error_response("File size must be under 10 MB"), status=400)

        # Save file
        seeker_dir = os.path.join(UPLOAD_DIR, "seekers", str(request.seeker.id))
        os.makedirs(seeker_dir, exist_ok=True)

        fname = f"{uuid.uuid4()}_{file.name}"
        file_path = os.path.join(seeker_dir, fname)
        with open(file_path, "wb+") as f:
            for chunk in file.chunks():
                f.write(chunk)

        # Parse with existing agent (async → sync)
        file_ext = os.path.splitext(file.name.lower())[1].lstrip(".")  # pdf, docx, txt etc.
        if file_ext == "doc":
            file_ext = "docx"
        parser = ResumeParsingAgent()
        from asgiref.sync import async_to_sync
        result = async_to_sync(parser.parse)(file_path, file_ext)
        parsed = result.get("parsed", {})

        # Normalize skills
        raw_skills = parsed.get("skills", [])

        def flatten_skill(s):
            """Convert skill object {skill, level, ...} or plain string to a string."""
            if isinstance(s, str):
                return s
            if isinstance(s, dict):
                return s.get("canonical_skill") or s.get("skill") or s.get("raw_skill") or s.get("name") or str(s)
            return str(s)

        # Flatten raw skills to strings first
        raw_skills_flat = [flatten_skill(s) for s in raw_skills if s]

        try:
            norm_agent = SkillNormalizationAgent()
            normalized_skills = async_to_sync(norm_agent.normalize)(raw_skills_flat)
            # Ensure normalized output is also flat strings
            normalized_skills = [flatten_skill(s) for s in normalized_skills if s]
        except Exception as norm_err:
            logger.warning("Skill normalization failed: %s", norm_err)
            normalized_skills = raw_skills_flat

        # Save to seeker account
        seeker = request.seeker
        seeker.resume_file_path = file_path
        seeker.resume_data = parsed
        seeker.skills = normalized_skills

        # Auto-fill headline if empty
        if not seeker.headline and parsed.get("current_role"):
            seeker.headline = parsed["current_role"]

        seeker.save(update_fields=["resume_file_path", "resume_data", "skills", "headline"])

        return JsonResponse(success_response({
            "message": "Resume uploaded and parsed successfully",
            "parsed": {
                "name": parsed.get("name"),
                "email": parsed.get("email"),
                "skills": normalized_skills[:20],  # first 20
                "total_skills": len(normalized_skills),
                "experience_years": parsed.get("total_experience_years", 0),
                "education": parsed.get("education", []),
                "summary": parsed.get("summary", ""),
            }
        }))

    except Exception as e:
        logger.error("Resume upload error: %s", e)
        return JsonResponse(error_response(f"Server error: {e}"), status=500)


@csrf_exempt
@require_seeker_jwt
def enhance_resume(request):
    """
    POST /api/v1/seeker/resume/enhance
    Run the AI Resume Enhancer on the seeker's current resume.
    Body (optional): { "job_description": "..." }
    """
    if request.method != "POST":
        return JsonResponse(error_response("Method not allowed"), status=405)
    try:
        seeker = request.seeker

        if not seeker.resume_data:
            return JsonResponse(error_response("Please upload a resume first"), status=400)

        import json as _json
        body = {}
        if request.body:
            try:
                body = _json.loads(request.body)
            except Exception:
                pass

        job_description = body.get("job_description", "")

        enhancer = ResumeEnhancerAgent()
        result = enhancer.enhance(seeker.resume_data, job_description)

        if result["success"]:
            # Store the enhanced version
            seeker.enhanced_resume = result["data"]
            seeker.save(update_fields=["enhanced_resume"])

        return JsonResponse(success_response(result["data"]))

    except Exception as e:
        logger.error("Resume enhance error: %s", e)
        return JsonResponse(error_response(f"Server error: {e}"), status=500)


@csrf_exempt
@require_seeker_jwt
def get_resume(request):
    """
    GET /api/v1/seeker/resume
    Returns the seeker's current parsed resume and enhanced version.
    """
    if request.method != "GET":
        return JsonResponse(error_response("Method not allowed"), status=405)
    seeker = request.seeker
    return JsonResponse(success_response({
        "has_resume": bool(seeker.resume_file_path or seeker.resume_data),
        "resume_data": seeker.resume_data,
        "enhanced_resume": seeker.enhanced_resume,
        "skills": seeker.skills,
    }))


@csrf_exempt
@require_seeker_jwt
def download_enhanced_resume_file(request):
    """
    POST /api/v1/seeker/resume/download
    Download the enhanced resume as .docx or .txt, or apply placeholder replacement to an uploaded template.
    """
    if request.method != "POST":
        return JsonResponse(error_response("Method not allowed"), status=405)
    
    try:
        seeker = request.seeker
        enhanced_resume = seeker.enhanced_resume
        resume_data = seeker.resume_data
        
        if not resume_data:
            return JsonResponse(error_response("Please upload and parse a resume first"), status=400)
            
        data_to_use = enhanced_resume if (enhanced_resume and enhanced_resume.get("enhanced_experience")) else resume_data
        
        import json as _json
        file_format = "docx"
        template_type = "modern"
        
        if request.content_type == "application/json" or "application/json" in request.headers.get("Content-Type", ""):
            body = _json.loads(request.body) if request.body else {}
            file_format = body.get("format", "docx")
            template_type = body.get("template_type", "modern")
        else:
            file_format = request.POST.get("format", "docx")
            template_type = request.POST.get("template_type", "modern")
            
        file_format = file_format.lower()
        template_type = template_type.lower()
        
        # TXT FORMAT
        if file_format == "txt":
            content = []
            content.append(f"{resume_data.get('name', seeker.full_name).upper()}")
            contact_info = []
            if resume_data.get("email"): contact_info.append(resume_data.get("email"))
            if seeker.phone: contact_info.append(seeker.phone)
            if seeker.location: contact_info.append(seeker.location)
            content.append(" | ".join(contact_info))
            content.append("\n" + "="*40 + "\n")
            
            summary_text = data_to_use.get("summary_rewrite") or resume_data.get("summary")
            if summary_text:
                content.append("PROFESSIONAL SUMMARY")
                content.append("-" * 20)
                content.append(summary_text + "\n")
                
            content.append("WORK EXPERIENCE")
            content.append("-" * 25)
            experience = data_to_use.get("enhanced_experience") or resume_data.get("experience") or []
            for exp in experience:
                if not isinstance(exp, dict):
                    content.append(str(exp))
                    content.append("")
                    continue
                role = exp.get("role") or exp.get("title") or "Role"
                company = exp.get("company") or "Company"
                content.append(f"{role} - {company}")
                bullets = exp.get("enhanced_bullets") or exp.get("bullets") or []
                for b in bullets:
                    content.append(f"  • {b}")
                content.append("")
                
            # PROJECTS
            projects = data_to_use.get("enhanced_projects") or resume_data.get("projects") or []
            if projects:
                content.append("PROJECTS")
                content.append("-" * 25)
                for proj in projects:
                    if not isinstance(proj, dict):
                        content.append(str(proj))
                        content.append("")
                        continue
                    name = proj.get("name") or proj.get("title") or "Project"
                    content.append(name)
                    bullets = proj.get("enhanced_bullets") or proj.get("bullets") or []
                    if bullets:
                        for b in bullets:
                            content.append(f"  • {b}")
                    else:
                        desc = proj.get("description") or ""
                        if isinstance(desc, str):
                            for line in desc.split("\n"):
                                if line.strip():
                                    content.append(f"  • {line.strip()}")
                    content.append("")
                
            content.append("TECHNICAL SKILLS")
            content.append("-" * 20)
            skills = data_to_use.get("skills") or resume_data.get("skills") or []
            skills_flat = [s if isinstance(s, str) else s.get("skill", "") for s in skills]
            content.append(", ".join(skills_flat) + "\n")
            
            education = resume_data.get("education") or []
            if education:
                content.append("EDUCATION")
                content.append("-" * 15)
                for edu in education:
                    deg = edu.get("degree") or "Degree"
                    inst = edu.get("institution") or "Institution"
                    content.append(f"{deg} - {inst}")
                    
            text_body = "\n".join(content)
            
            from django.http import HttpResponse
            response = HttpResponse(text_body, content_type="text/plain")
            response["Content-Disposition"] = 'attachment; filename="Enhanced_Resume.txt"'
            return response
            
        # DOCX FORMAT
        elif file_format == "docx":
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import io
            
            custom_file = request.FILES.get("template_file")
            
            if template_type == "custom" and custom_file:
                doc = Document(custom_file)
                flat_skills = [s if isinstance(s, str) else s.get("skill", "") for s in resume_data.get("skills", [])]
                
                exp_list = []
                experience = data_to_use.get("enhanced_experience") or resume_data.get("experience") or []
                for exp in experience:
                    if not isinstance(exp, dict):
                        exp_list.append(str(exp))
                        continue
                    role = exp.get("role") or exp.get("title") or "Role"
                    company = exp.get("company") or "Company"
                    bullets = "\n".join([f"• {b}" for b in (exp.get("enhanced_bullets") or exp.get("bullets") or [])])
                    exp_list.append(f"{role} @ {company}\n{bullets}")
                    
                proj_list = []
                projects = data_to_use.get("enhanced_projects") or resume_data.get("projects") or []
                for proj in projects:
                    if not isinstance(proj, dict):
                        proj_list.append(str(proj))
                        continue
                    name = proj.get("name") or proj.get("title") or "Project"
                    bullets = proj.get("enhanced_bullets") or proj.get("bullets") or []
                    if bullets:
                        bullets_str = "\n".join([f"• {b}" for b in bullets])
                    else:
                        desc = proj.get("description") or ""
                        if isinstance(desc, str):
                            bullets_str = "\n".join([f"• {line.strip()}" for line in desc.split("\n") if line.strip()])
                        else:
                            bullets_str = str(desc)
                    proj_list.append(f"{name}\n{bullets_str}")
                    
                edu_list = []
                education = resume_data.get("education") or []
                for edu in education:
                    edu_list.append(f"{edu.get('degree')} from {edu.get('institution')}")
                
                replacements = {
                    "{{full_name}}": resume_data.get("name", seeker.full_name),
                    "{{email}}": resume_data.get("email", seeker.email or ""),
                    "{{phone}}": seeker.phone or "",
                    "{{location}}": seeker.location or "",
                    "{{headline}}": seeker.headline or "",
                    "{{summary}}": data_to_use.get("summary_rewrite") or resume_data.get("summary") or "",
                    "{{skills}}": ", ".join(flat_skills),
                    "{{experience}}": "\n\n".join(exp_list),
                    "{{projects}}": "\n\n".join(proj_list),
                    "{{education}}": "\n".join(edu_list)
                }
                
                for p in doc.paragraphs:
                    for key, val in replacements.items():
                        if key in p.text:
                            for run in p.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(val))
                                    
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for key, val in replacements.items():
                                    if key in p.text:
                                        for run in p.runs:
                                            if key in run.text:
                                                run.text = run.text.replace(key, str(val))
            else:
                doc = Document()
                for s in doc.sections:
                    s.top_margin = Inches(0.8)
                    s.bottom_margin = Inches(0.8)
                    s.left_margin = Inches(0.8)
                    s.right_margin = Inches(0.8)
                    
                primary_color = RGBColor(17, 17, 17)
                if template_type == "tech":
                    primary_color = RGBColor(79, 70, 229)
                elif template_type == "minimal":
                    primary_color = RGBColor(75, 85, 99)
                    
                name_p = doc.add_paragraph()
                name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                name_run = name_p.add_run(resume_data.get("name", seeker.full_name).upper())
                name_run.bold = True
                name_run.font.size = Pt(20)
                name_run.font.color.rgb = primary_color
                
                contact_p = doc.add_paragraph()
                contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_parts = []
                if resume_data.get("email"): contact_parts.append(resume_data.get("email"))
                if seeker.phone: contact_parts.append(seeker.phone)
                if seeker.location: contact_parts.append(seeker.location)
                contact_run = contact_p.add_run("  |  ".join(contact_parts))
                contact_run.font.size = Pt(9.5)
                
                summary_text = data_to_use.get("summary_rewrite") or resume_data.get("summary")
                if summary_text:
                    h = doc.add_heading("PROFESSIONAL SUMMARY", level=2)
                    for run in h.runs:
                        run.font.color.rgb = primary_color
                        run.font.size = Pt(12)
                    p = doc.add_paragraph(summary_text)
                    p.paragraph_format.space_after = Pt(10)
                    
                h = doc.add_heading("WORK EXPERIENCE", level=2)
                for run in h.runs:
                    run.font.color.rgb = primary_color
                    run.font.size = Pt(12)
                experience = data_to_use.get("enhanced_experience") or resume_data.get("experience") or []
                for exp in experience:
                    if not isinstance(exp, dict):
                        doc.add_paragraph(str(exp))
                        continue
                    p = doc.add_paragraph()
                    r_run = p.add_run(exp.get("role") or exp.get("title") or "Role")
                    r_run.bold = True
                    c_run = p.add_run(f"  -  {exp.get('company') or 'Company'}")
                    c_run.italic = True
                    bullets = exp.get("enhanced_bullets") or exp.get("bullets") or []
                    for b in bullets:
                        doc.add_paragraph(b, style='List Bullet')
                        
                projects = data_to_use.get("enhanced_projects") or resume_data.get("projects") or []
                if projects:
                    h = doc.add_heading("PROJECTS", level=2)
                    for run in h.runs:
                        run.font.color.rgb = primary_color
                        run.font.size = Pt(12)
                    for proj in projects:
                        if not isinstance(proj, dict):
                            doc.add_paragraph(str(proj))
                            continue
                        p = doc.add_paragraph()
                        p_run = p.add_run(proj.get("name") or proj.get("title") or "Project")
                        p_run.bold = True
                        bullets = proj.get("enhanced_bullets") or proj.get("bullets") or []
                        if bullets:
                            for b in bullets:
                                doc.add_paragraph(b, style='List Bullet')
                        else:
                            desc = proj.get("description") or ""
                            if isinstance(desc, str):
                                for line in desc.split("\n"):
                                    if line.strip():
                                        doc.add_paragraph(line.strip(), style='List Bullet')
                        
                h = doc.add_heading("TECHNICAL SKILLS", level=2)
                for run in h.runs:
                    run.font.color.rgb = primary_color
                    run.font.size = Pt(12)
                skills = data_to_use.get("skills") or resume_data.get("skills") or []
                skills_flat = [s if isinstance(s, str) else s.get("skill", "") for s in skills]
                doc.add_paragraph(", ".join(skills_flat))
                
                education = resume_data.get("education") or []
                if education:
                    h = doc.add_heading("EDUCATION", level=2)
                    for run in h.runs:
                        run.font.color.rgb = primary_color
                        run.font.size = Pt(12)
                    for edu in education:
                        p = doc.add_paragraph()
                        d_run = p.add_run(edu.get("degree") or "Degree")
                        d_run.bold = True
                        i_run = p.add_run(f"  -  {edu.get('institution') or 'Institution'}")
                        i_run.italic = True
                        
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            
            from django.http import HttpResponse
            response = HttpResponse(file_stream.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            response["Content-Disposition"] = 'attachment; filename="Enhanced_Resume.docx"'
            return response
            
        return JsonResponse(error_response("Unsupported download format"), status=400)
        
    except Exception as e:
        logger.error("Download enhanced resume error: %s", e)
        return JsonResponse(error_response(f"Server error: {e}"), status=500)


@csrf_exempt
@require_seeker_jwt
def check_ats_score(request):
    """
    POST /api/v1/seeker/resume/check-ats
    Run the official ATS checker against a job description.
    """
    if request.method != "POST":
        return JsonResponse(error_response("Method not allowed"), status=405)
        
    try:
        seeker = request.seeker
        enhanced_resume = seeker.enhanced_resume
        resume_data = seeker.resume_data
        
        if not resume_data:
            return JsonResponse(error_response("Please upload and parse a resume first"), status=400)
            
        import json as _json
        body = {}
        if request.body:
            try:
                body = _json.loads(request.body)
            except Exception:
                pass
                
        job_description = body.get("job_description", "")
        
        # If enhanced resume exists and contains enhanced experience, form virtual enhanced resume:
        if enhanced_resume and enhanced_resume.get("enhanced_experience"):
            active_resume_data = {
                "personalInfo": resume_data.get("personalInfo", {}),
                "summary": enhanced_resume.get("professional_summary_enhanced") or enhanced_resume.get("summary_rewrite") or resume_data.get("summary") or resume_data.get("professional_summary") or "",
                "skills": list(set(resume_data.get("skills", []) + enhanced_resume.get("missing_keywords", []))),
                "experience": [],
                "education": resume_data.get("education", []),
                "projects": [],
                "certifications": resume_data.get("certifications", []),
                "languages": resume_data.get("languages", [])
            }
            
            # Map enhanced experiences
            for exp in resume_data.get("experience", []):
                company = exp.get("company", "")
                role = exp.get("title", "") or exp.get("role", "")
                bullets = exp.get("bullets", []) or exp.get("responsibilities", []) or []
                
                # Check for enhanced bullets match
                for ee in enhanced_resume.get("enhanced_experience", []):
                    if ee.get("company", "").lower() == company.lower() or ee.get("role", "").lower() == role.lower():
                        bullets = ee.get("enhanced_bullets", bullets)
                        break
                active_resume_data["experience"].append({
                    **exp,
                    "bullets": bullets
                })

            # Map enhanced projects
            for proj in resume_data.get("projects", []):
                name = proj.get("name", "")
                desc = proj.get("description", "")
                
                # Check for enhanced project description match
                for ep in enhanced_resume.get("enhanced_projects", []):
                    if ep.get("name", "").lower() == name.lower():
                        eb = ep.get("enhanced_bullets", [])
                        if eb:
                            desc = "\n".join(eb)
                        else:
                            desc = ep.get("enhanced_description", desc)
                        break
                active_resume_data["projects"].append({
                    **proj,
                    "description": desc
                })
        else:
            active_resume_data = resume_data

        # Use the deterministic AtsCompatibilityAgent
        from agents.ats_compatibility_agent import AtsCompatibilityAgent
        ats_agent = AtsCompatibilityAgent()
        report = ats_agent.analyze(None, active_resume_data, job_description)
        
        score = report.get("overallScore", 0)
        
        # Calculate verdict
        if score >= 90:
            verdict = "Excellent Match"
        elif score >= 80:
            verdict = "Good Match"
        elif score >= 60:
            verdict = "Fair Match"
        else:
            verdict = "Poor Match"
            
        bd = report.get("detailed_breakdown", {})
        fmt_score = bd.get("ats_formatting", {}).get("score", 100)
        fmt_issues = bd.get("ats_formatting", {}).get("issues", [])
        
        kw_score = bd.get("keyword_match", {}).get("score", 100)
        matched_kws = bd.get("keyword_match", {}).get("matched", [])
        
        edu_score = bd.get("education_match", {}).get("score", 100)
        
        # Build standard output checks matching frontend expectations
        if fmt_score >= 90:
            format_check = "Passed (Clean headers & margins)"
        else:
            format_check = f"Needs Improvement ({len(fmt_issues)} formatting warnings)"
            
        if kw_score >= 80:
            keyword_check = "Passed (Critical keywords integrated)"
        else:
            keyword_check = "Needs Improvement (Missing target keywords)"
            
        if edu_score >= 90:
            structure_check = "Passed (Standard sections & credentials detected)"
        else:
            structure_check = "Needs Improvement (Degree match warning)"
            
        result_json = {
            "score": score,
            "verdict": verdict,
            "matched_keywords": matched_kws,
            "format_check": format_check,
            "keyword_check": keyword_check,
            "structure_check": structure_check,
            "recommendations": report.get("topSuggestions", [])
        }
        
        return JsonResponse(success_response(result_json))
        
    except Exception as e:
        logger.error("ATS checking error: %s", e)
        return JsonResponse(error_response(f"Server error: {e}"), status=500)
